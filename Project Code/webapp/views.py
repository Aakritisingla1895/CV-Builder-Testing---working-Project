from unicodedata import name
from django.shortcuts import render, redirect
from django.contrib.auth.forms import PasswordResetForm
from .models import StudentUser, PersonalInformationQuestionResponse,WorkExperienceQuestionResponse, TechnicalSkillsQuestionResponse, CertificationsQuestionsResponse, ProfessionalSummaryQuestionResponse, EducationQuestionResponse
from .forms import PDFUploadForm
from django.contrib.sessions.models import Session
from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from django.contrib.auth import authenticate, login, logout
from docx.shared import Inches 
from spacy.language import Language
from spacy_langdetect import LanguageDetector # spacy library - function to detect language of a text
import pandas
from pandas import *
import docx # library to auto-genearte word doc
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from django.http import  HttpResponse, FileResponse
import re
import openai
# import the generated API key from the secret_key file
from .secret_key import API_KEY
# loading the API key from the secret_key file
openai.api_key = API_KEY

from .forms import PDFUploadForm, Page1Form, Page2Form, Page3Form, Page4Form, Page5Form,Page6Form

import os

keywords = ['Gap Analysis','Process Mapping','Prototyping','UML Diagram','User Stories','Agile Frameworks','Requirement Gathering','Full Project Lifecycle',
'Workshop Facilitation','Target Operating Model','Microsoft Visio','Microsoft Project','Draw IO','Asana','Basecamp','Balsamiq','Agile Methodologies',
'Waterfall Methodologies','Business Process Improvement','Business Analysis','Requirements Analysis','Digital & Tech Transformation','Business Transformation',
'Performance Management','AS IS and TO BE','Scrum meetings','Facilitating meetings and workshops','interviewing stakeholders']

transferable_skills = [
    "Communication Skills","Written communication","Verbal communication","Analytical Skills","Critical thinking","Problem-solving",
    "Data analysis","Technical Proficiency","Business analysis tools","Microsoft Excel","Visio","Project management software",
    "Basic understanding of databases","Project Management","Planning and organization","Stakeholder management",
    "Requirements Gathering","Elicitation techniques","Documenting requirements","Process Modeling","Business process modeling",
    "BPMN (Business Process Model and Notation)","Negotiation Skills","Resolving conflicts","Adaptability","Flexibility",
    "Continuous learning","Attention to Detail","Accuracy","Decision-Making","Evaluating options","Risk analysis",
    "Interpersonal Skills","Collaboration","Empathy","Time Management","Prioritization","Presentation Skills","Delivering presentations",
    "Quality Assurance","Testing and validation","Leadership Skills","Leading meetings","Influence"
]

def home(request):

    return render(request,'home.html')

def login_screen(request):
    
    return render(request, 'index.html')

def signup(request):
    if request.method == 'POST':
        name = request.POST['name']
        email = request.POST['email']
        password = request.POST['password']
        confirm_password = request.POST['confirm_password']

        if password == confirm_password:
            user = StudentUser.objects.create_user(name=name, email=email, password=password)
            user.save()
            login(request, user)
            return redirect('/')
        else:
            error_message = 'Passwords must match'
    return render(request, 'index.html')



def userlogin(request):
   
    if request.method == 'POST':
        name = request.POST['name']
        password = request.POST['password']

        user = authenticate(name=name, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')
        else:
            error_message = 'Invalid username or password'

        
    return render(request,'index.html')


def logout_view(request):
    logout(request)
    return redirect('/')

def home(request):

    return render(request,'home.html')

def about_us(request):
    
    return render(request,'about-us.html')

def uploadPdf(request):
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            title = form.cleaned_data['title']
            pdf_file = form.cleaned_data['pdf_file']

            # If you have a model, you can save the PDF like this:
            # pdf_document = PDFDocument(title=title, pdf_file=pdf_file)
            # pdf_document.save()

            # If you just want to store the file without a model, you can save it to a specific directory
            # with a unique filename.
            # with open('path/to/your/pdf/files/' + pdf_file.name, 'wb') as destination:
            #     for chunk in pdf_file.chunks():
            #         destination.write(chunk)

            return redirect('success')
    else:
        form = PDFUploadForm()

    return render(request, 'home.html', {'form': form})

def processingCV(request):
    
    return render(request, 'cv_file.html')


def question_form(request):
    if request.method == 'POST':
        form = Page1Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('professional_experience_form')  # Redirect to the next page
    else:
        form = Page1Form()


    return render(request, 'personal_information_form.html', {'form': form})

def professional_experience_form(request):
    if request.method == 'POST':
        form = Page2Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('work_experience_form')  # Redirect to the thank you page or another page
    else:
        form = Page2Form()
    return render(request, 'professional_experience_form.html', {'form': form})

def work_experience_form(request):
    if request.method == 'POST':
        form = Page3Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('technical_skills_form')  # Redirect to the thank you page or another page
    else:
        form = Page3Form()

    return render(request, 'work_experience_form.html', {'form': form})

def technical_skills_form(request):
    if request.method == 'POST':
        form = Page4Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('certifications_form')  # Redirect to the thank you page or another page
    else:
        form = Page4Form()

    return render(request, 'technical_skills_form.html', {'form': form})


def certifications_form(request):
    if request.method == 'POST':
        form = Page5Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('education_form')  # Redirect to the thank you page or another page
    else:
        form = Page5Form()

    return render(request, 'certifications_form.html', {'form': form})



def education_form(request):
    if request.method == 'POST':
        form = Page6Form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('chatgptBackend')  # Redirect to the thank you page or another page
    else:
        form = Page6Form()

    return render(request, 'education_form.html', {'form': form})


def progress_bar(request):
    
    return render(request, 'progress_bar.html')



def chatgptBackend(request):
    
    # Define the system message
    system_message = {"role": "system", "content": "You are an intelligent assistant."}

    # Initialize messages with the system message
    messages = [system_message]
    
     # Retrieve the latest user response from the database
    latest_response = PersonalInformationQuestionResponse.objects.last()
    work_experience_response= WorkExperienceQuestionResponse.objects.last()
    professional_experience_response = ProfessionalSummaryQuestionResponse.objects.last()
    technical_skills_response = TechnicalSkillsQuestionResponse.objects.last()
    education_response = EducationQuestionResponse.objects.last()
    

    # Define user prompts
    prompt1 = "Give a professional summary for a Business Analyst profile in 4 sentences summarizing below content"
    prompt2 = f"Give a summary of {work_experience_response.question_10} and {work_experience_response.question_12}"
    prompt3 = f"List down in 6 bullet points picking from {technical_skills_response.question_15} and {technical_skills_response.question_16} and pick only 6 imp words from this , highlighting them as skills"
    prompt4 = f"can you pick up names of places of any college name , university name and major areas of study from {education_response.question_20} {education_response.question_21}"
    
    

    # Process prompt1
    if latest_response:  # Replace with your condition
        question_1_data = latest_response.question_1
        question_2_data = latest_response.question_2
        messages.append({"role": "user", "content": prompt1})
        chat_response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
        assistant_reply1 = chat_response.choices[0].message.content
        print(f"User Prompt 1: {prompt1}")
        print(f"ChatGPT Response 1: {assistant_reply1}\n")
        
        
    # Process prompt2
    if work_experience_response:  # Replace with your condition
        work_experience_response_data_1 = work_experience_response.question_10
        messages.append({"role": "user", "content": prompt2})
        chat_response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
        assistant_reply2 = chat_response.choices[0].message.content
        print(f"User Prompt 2: {prompt2}")
        print(f"ChatGPT Response 2: {assistant_reply2}\n")
        
    # Process prompt2
    if technical_skills_response:  # Replace with your condition
       
        messages.append({"role": "user", "content": prompt2})
        chat_response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
        assistant_reply3 = chat_response.choices[0].message.content
        print(f"User Prompt 3: {prompt3}")
        print(f"ChatGPT Response 3: {assistant_reply3}\n")
        
        
    if education_response:  # Replace with your condition
       
        messages.append({"role": "user", "content": prompt2})
        chat_response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
        assistant_reply4 = chat_response.choices[0].message.content
        print(f"User Prompt 4: {prompt4}")
        print(f"ChatGPT Response 4: {assistant_reply4}\n")

        
    context = {'reply':assistant_reply1, 'latest_response':latest_response, 'question_1_data':question_1_data, 'work_experience_response':work_experience_response, 'assistant_reply2':assistant_reply2, 
               'assistant_reply3':assistant_reply3, 'technical_skills_response':technical_skills_response,
               'education_response':education_response, 'assistant_reply4':assistant_reply4}
    return render(request, 'build_cv.html',context)
    
def create_document(request):
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run('\t\t\t\t\t')
    r.add_picture('MicrosoftTeams-image.png')
    
    return render('show_buildcv.html')

def generate_word_template(request):
    context = {
        'name': 'John Doe',
        'title': 'Business Analyst',
        'contact_info': [
            {'label': 'Location Preferences:', 'value': 'Anywhere'},
            {'label': 'Availability:', 'value': 'Immediate'},
            {'label': 'LinkedIn Profile:', 'value': 'www.linkedin.com/in/johndoe'},
            {'label': 'Phone:', 'value': '+1-234-567-8901'},
            {'label': 'Email:', 'value': 'johndoe@example.com'},
        ],
        'profile': 'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Mauris at mauris vitae eros tristique semper. Donec sed quam nec mauris tincidunt semper.',
        'education': [
            {'institution': 'Stanford University', 'location': 'Palo Alto, CA', 'graduation_date': '2020'},
        ],
        'experience': [
            {'company': 'Google', 'location': 'Mountain View, CA', 'dates_employed': '2020-Present', 'responsibilities': ['Software Engineer', 'Developed web applications using Python and JavaScript']},
        ],
    }
    
    return render(request, 'test_cv.html',context)



def test_front(request):
    
    return render(request, 'test_cv.html')


def test_buildCV(request):
    messages = [ {"role": "system", "content":  
              "You are a intelligent assistant."} ] 
    
    
    userQAresponse = "I have a solid background in IT with 3 years of experience. I have worked in various roles, including software engineer, where I've gained expertise in business analysis. My professional journey has been characterized by a commitment to  and a focus on delivering high-quality results. I am proficient in a variety of tools commonly used for business analysis. This includes advanced proficiency in tools such as Microsoft Excel for data analysis and visualization, SQL for database querying, and Tableau for creating insightful dashboards. Additionally, I have experience with business intelligence tools like Power BI and statistical analysis tools like R and Python. These tools have allowed me to extract meaningful insights from complex datasets, facilitating informed decision-making within the organizations I've worked for.One of my notable achievements is [provide a specific achievement]. In my previous role at [company], I spearheaded a [mention project or initiative] that resulted in [quantifiable result, such as increased efficiency, cost savings, revenue growth, etc.]. This accomplishment not only showcased my ability to [highlight a key skill or competency] but also had a tangible impact on the overall success of the team and the organization."
    
    # Retrieve the latest user response from the database
    latest_response = PersonalInformationQuestionResponse.objects.last()

    # Access the data from the retrieved response
    if latest_response:
        question_1_data = latest_response.question_1
        # Add your logic here
        print(f"Retrieved data in chatgpt function: {question_1_data}")
        
        message = "Give a professional summary for Business Analysts profile in 4 sentences summarizing below content"
        if message: 
            messages.append( 
                {"role": "user", "content": message}, 
            ) 
            chat = openai.ChatCompletion.create( 
                model="gpt-3.5-turbo", messages=messages 
            ) 
        reply = chat.choices[0].message.content 
        print(f"ChatGPT: {reply}") 
        messages.append({"role": "assistant", "content": reply}) 
        
         # Pass the data to the template context
       
        context = {'reply':reply, 'question_1_data': question_1_data}
    
        return render(request, 'build_cv_template.html', context)
    
    
def new_file(request):
    
    return render(request, 'new_file.html')
    
    
def how_to_use(request):
    
    return render(request, 'how_to_use.html')


def templates(request):
    
    return render(request, 'templates.html')

def template_1(request):
    return render(request, 'template_1.html')


def template_2(request):
    return render(request, 'template_2.html')


def qa_animation(request):
    
    return render(request, 'qa_animation.html')